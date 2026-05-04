"""
template_filler.py
------------------
Fills the Pro Talent master template with extracted candidate data.

Handles:
- Simple {{placeholder}} replacements in paragraphs and tables
- Multi-bullet expansion for duties, achievements, and computer skills (one bullet per item)
- Empty slots (e.g. only 2 jobs but template has 3) — fills extras with '[Not applicable]'
"""

from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# Generic placeholder replacement (preserves formatting)
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Generic placeholder replacement (preserves formatting at the run level)
# ---------------------------------------------------------------------------

def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> bool:
    """
    Replace {{placeholder}} tokens in a paragraph while preserving each run's
    original formatting (font, size, bold, italic, colour).

    The challenge: placeholders can be split across multiple runs by Word
    (e.g. {{ in run 1, surname}} in run 2). We solve this by:
      1. Concatenating run texts into a single string
      2. Doing the replacement
      3. Distributing the result back across runs based on each run's
         original character-length proportion
    But for the simpler (much more common) case where a placeholder lives
    entirely within one run, we replace it in-place — preserving that run's
    formatting exactly.
    """
    if not paragraph.runs:
        return False

    full_text = paragraph.text
    if not any(key in full_text for key in replacements):
        return False

    # Strategy: walk the runs and for each run, see if it contains a placeholder
    # entirely within itself. If so, replace in place — preserves formatting.
    # If a placeholder is split across runs, we fall back to the more complex
    # whole-paragraph rewrite (which sacrifices intra-paragraph formatting).

    any_replacement = False
    for run in paragraph.runs:
        run_text = run.text
        new_text = run_text
        for key, value in replacements.items():
            if key in new_text:
                new_text = new_text.replace(key, str(value))
        if new_text != run_text:
            run.text = new_text
            any_replacement = True

    # If after the per-run replacement there are still tokens in the paragraph,
    # they must be split across runs. Fall back to the original "collapse into
    # first run" approach for those — but only as a last resort.
    remaining_text = paragraph.text
    if any(key in remaining_text for key in replacements):
        new_text = remaining_text
        for key, value in replacements.items():
            new_text = new_text.replace(key, str(value))
        # Put the rebuilt text into the first non-empty run to preserve at
        # least one run's formatting; clear the others.
        first_run = paragraph.runs[0]
        first_run.text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
        any_replacement = True

    return any_replacement


def _replace_in_table(table, replacements: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            for nested in cell.tables:
                _replace_in_table(nested, replacements)


def _replace_everywhere(doc: DocxDocument, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        _replace_in_table(table, replacements)
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in header_footer.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            for table in header_footer.tables:
                _replace_in_table(table, replacements)


# ---------------------------------------------------------------------------
# Multi-bullet expansion
# ---------------------------------------------------------------------------

def _find_paragraph_with_token(doc: DocxDocument, token: str):
    """Find the first paragraph containing token. Returns Paragraph or None."""
    for p in doc.paragraphs:
        if token in p.text:
            return p
    return None


def _insert_paragraph_after(paragraph, text: str) -> Paragraph:
    """
    Insert a new paragraph after the given one, copying its formatting (including bullet style).
    """
    new_p = deepcopy(paragraph._element)

    # Clear runs in the copy
    for r in new_p.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"):
        new_p.remove(r)

    # Insert after the original paragraph
    paragraph._element.addnext(new_p)

    # Wrap in Paragraph object and add the text
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def _expand_bulleted_token(doc: DocxDocument, token: str, items: list[str]) -> None:
    """
    Replace a token like {{job1_duties}} with a series of bulleted paragraphs.
    The paragraph containing the token must already be in the right bullet style
    (which it is, since we put the token where the original "M" bullet was).

    items: list of strings (one per bullet point)

    Also removes the trailing empty bullet that follows the placeholder paragraph
    (left over from the original template structure).
    """
    paragraph = _find_paragraph_with_token(doc, token)
    if paragraph is None:
        return

    # Find the next paragraph — if it's an empty bullet (leftover from template), mark for removal.
    # We use the body element ordering since paragraph._element.getnext() walks XML siblings.
    next_element = paragraph._element.getnext()
    next_para_to_remove = None
    if next_element is not None and next_element.tag.endswith("}p"):
        next_para = Paragraph(next_element, paragraph._parent)
        if not next_para.text.strip():
            # Likely the empty bullet — flag for removal AFTER we insert new bullets
            next_para_to_remove = next_para

    if not items:
        # No items — replace token with placeholder text
        _replace_in_paragraph(paragraph, {token: "(info absent on CV)"})
    else:
        # First item goes in the existing paragraph (preserves bullet styling)
        _replace_in_paragraph(paragraph, {token: items[0]})

        # Remaining items get inserted as new paragraphs after the original
        last_inserted = paragraph
        for item in items[1:]:
            last_inserted = _insert_paragraph_after(last_inserted, item)

    # Remove the leftover empty bullet
    if next_para_to_remove is not None:
        try:
            elem = next_para_to_remove._element
            elem.getparent().remove(elem)
        except Exception:
            pass


def _expand_text_block_token(doc: DocxDocument, token: str, lines: list[str]) -> None:
    """
    Replace a token in a paragraph with a multi-line text block.
    The first line replaces the token (preserves the paragraph's formatting),
    subsequent lines get inserted as new paragraphs after.

    Unlike _expand_bulleted_token, this doesn't try to remove a trailing empty
    paragraph — it's used for non-bulleted blocks like references where the
    paragraph stands alone.

    lines: list of strings (one per line)
    """
    paragraph = _find_paragraph_with_token(doc, token)
    if paragraph is None:
        return

    if not lines:
        _replace_in_paragraph(paragraph, {token: "(info absent on CV)"})
        return

    _replace_in_paragraph(paragraph, {token: lines[0]})

    last_inserted = paragraph
    for line in lines[1:]:
        last_inserted = _insert_paragraph_after(last_inserted, line)


def _format_reference_lines(reference: dict) -> list[str]:
    """
    Format a single reference dict as a list of display lines.
    Returns lines in order: Name, Company, Phone.
    Skips fields that are empty / missing entirely (so a reference with no
    company doesn't show a stray "Company:" line).
    """
    if not isinstance(reference, dict):
        return []

    lines = []
    name = reference.get("name", "").strip() if reference.get("name") else ""
    company = reference.get("company", "").strip() if reference.get("company") else ""
    phone = reference.get("phone", "").strip() if reference.get("phone") else ""

    if name:
        lines.append(f"Name: {name}")
    if company:
        lines.append(f"Company: {company}")
    if phone:
        lines.append(f"Phone: {phone}")

    return lines


# ---------------------------------------------------------------------------
# Data formatters
# ---------------------------------------------------------------------------

def _safe_str(value, default: str = "(info absent on CV)") -> str:
    if value is None or value == "" or value == "Not specified":
        return default
    return str(value)


def _list_or_default(value) -> list:
    """Return a list of items, filtering out falsy values but preserving types (dicts stay dicts)."""
    if isinstance(value, list):
        return [v for v in value if v]
    return []


def _str_list(value) -> list[str]:
    """Coerce to list of strings (used for plain string lists like achievements/skills)."""
    if isinstance(value, list):
        return [str(v) for v in value if v]
    return []


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def fill_template(
    template_path: str | Path,
    candidate_data: dict,
    output_path: str | Path,
) -> Path:
    """Fill the Pro Talent template with extracted candidate data."""

    doc = Document(template_path)

    # ---- Personal details (simple replacements) ----
    first_name = _safe_str(candidate_data.get("first_name"))
    surname = _safe_str(candidate_data.get("surname"))

    simple_replacements = {
        "{{candidate_first_name}}": first_name,
        "{{first_name}}": first_name,
        "{{surname}}": surname,
        "{{identity_number}}": _safe_str(candidate_data.get("identity_number")),
        "{{equity}}": _safe_str(candidate_data.get("equity")),
        "{{residential_area}}": _safe_str(candidate_data.get("residential_area")),
        "{{language}}": _safe_str(candidate_data.get("language")),
        "{{transport}}": _safe_str(candidate_data.get("transport")),
        "{{drivers_licence}}": _safe_str(candidate_data.get("drivers_licence")),
        "{{current_salary}}": _safe_str(candidate_data.get("current_salary")),
        "{{required_salary}}": _safe_str(candidate_data.get("required_salary")),
        "{{availability}}": _safe_str(candidate_data.get("availability")),
        "{{itc_criminal_record}}": "None",  # default — recruiter updates after interview
        "{{itc_reason}}": "(info absent on CV)",
    }

    # ---- Education (2 slots) ----
    education = _list_or_default(candidate_data.get("education"))
    for slot in range(1, 3):
        if slot <= len(education) and isinstance(education[slot - 1], dict):
            edu = education[slot - 1]
            simple_replacements[f"{{{{edu{slot}_institution}}}}"] = _safe_str(edu.get("institution"))
            simple_replacements[f"{{{{edu{slot}_date}}}}"] = _safe_str(edu.get("date"))
            simple_replacements[f"{{{{edu{slot}_qualification}}}}"] = _safe_str(edu.get("qualification"))
        else:
            simple_replacements[f"{{{{edu{slot}_institution}}}}"] = ""
            simple_replacements[f"{{{{edu{slot}_date}}}}"] = ""
            simple_replacements[f"{{{{edu{slot}_qualification}}}}"] = ""

    # ---- Employment history (3 slots) — simple fields ----
    jobs = _list_or_default(candidate_data.get("employment_history"))
    for slot in range(1, 4):
        if slot <= len(jobs) and isinstance(jobs[slot - 1], dict):
            job = jobs[slot - 1]
            simple_replacements[f"{{{{job{slot}_company}}}}"] = _safe_str(job.get("company"))
            simple_replacements[f"{{{{job{slot}_period}}}}"] = _safe_str(job.get("period"))
            simple_replacements[f"{{{{job{slot}_position}}}}"] = _safe_str(job.get("position"))
            simple_replacements[f"{{{{job{slot}_reason_for_leaving}}}}"] = _safe_str(
                job.get("reason_for_leaving"),
                default="(info absent on CV)"
            )
        else:
            # Empty slot
            simple_replacements[f"{{{{job{slot}_company}}}}"] = ""
            simple_replacements[f"{{{{job{slot}_period}}}}"] = ""
            simple_replacements[f"{{{{job{slot}_position}}}}"] = ""
            simple_replacements[f"{{{{job{slot}_reason_for_leaving}}}}"] = ""

    # ---- Apply all simple replacements ----
    _replace_everywhere(doc, simple_replacements)

    # ---- Multi-bullet expansions ----
    # Achievements (list of strings)
    achievements = _str_list(candidate_data.get("achievements"))
    _expand_bulleted_token(doc, "{{achievements_block}}", achievements)

    # Computer skills (list of strings)
    computer_skills = _str_list(candidate_data.get("computer_skills"))
    _expand_bulleted_token(doc, "{{computer_skills_block}}", computer_skills)

    # Duties for each job (list of strings inside each job dict)
    for slot in range(1, 4):
        if slot <= len(jobs) and isinstance(jobs[slot - 1], dict):
            duties = _str_list(jobs[slot - 1].get("duties"))
        else:
            duties = []
        _expand_bulleted_token(doc, f"{{{{job{slot}_duties}}}}", duties)

    # ---- References (2 slots) ----
    references = _list_or_default(candidate_data.get("references"))
    for slot in range(1, 3):
        if slot <= len(references) and isinstance(references[slot - 1], dict):
            ref_lines = _format_reference_lines(references[slot - 1])
        else:
            ref_lines = []
        _expand_text_block_token(doc, f"{{{{reference{slot}_block}}}}", ref_lines)

    # ---- Save ----
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def make_safe_filename(name: str) -> str:
    """Turn a candidate name into a safe filename component."""
    safe = "".join(c if c.isalnum() or c in "-_ " else "_" for c in name)
    return safe.strip().replace(" ", "_") or "candidate"


if __name__ == "__main__":
    print("This module is imported by app.py — run that instead.")
