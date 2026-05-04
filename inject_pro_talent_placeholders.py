"""
Inject placeholder tokens into the Pro Talent master template.
Handles the fact that section headings are inside tables (not regular paragraphs).
"""

from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph


def append_text_to_paragraph(paragraph, text: str) -> None:
    """Add a new run with text to the end of an existing paragraph."""
    run = paragraph.add_run(text)
    if len(paragraph.runs) > 1:
        prev = paragraph.runs[-2]
        if prev.font.name:
            run.font.name = prev.font.name
        if prev.font.size:
            run.font.size = prev.font.size


def replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace all text in a paragraph keeping the first run's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def find_paragraphs_between_headings(doc, after_heading: str, before_heading: str):
    """
    Walk through body elements. Return paragraphs that come AFTER the table containing
    `after_heading` and BEFORE the table containing `before_heading`.
    """
    body = doc.element.body
    found_first = False
    paragraphs = []

    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "tbl":
            table_text = "".join(child.itertext())
            if found_first and before_heading in table_text:
                break
            if after_heading in table_text:
                found_first = True
        elif tag == "p" and found_first:
            paragraphs.append(Paragraph(child, doc))
    return paragraphs


def remove_paragraph(paragraph) -> None:
    """Remove a paragraph from the document entirely."""
    p_element = paragraph._element
    p_element.getparent().remove(p_element)


def inject_placeholders(input_path: Path, output_path: Path) -> None:
    doc = Document(input_path)

    field_appends = {
        "SURNAME\t:": "{{surname}}",
        "FIRST NAME\t:": "{{first_name}}",
        "IDENTITY NUMBER\t:": "{{identity_number}}",
        "EQUITY\t:": "{{equity}}",
        "RESIDENTIAL AREA\t:": "{{residential_area}}",
        "LANGUAGE\t:": "{{language}}",
        "TRANSPORT\t:": "{{transport}}",
        "DRIVERS LICENCE\t:": "{{drivers_licence}}",
        "CURRENT SALARY\t:": "{{current_salary}}",
        "REQUIRED SALARY\t:": "{{required_salary}}",
        "AVAILABILITY\t:": "{{availability}}",
    }

    # 1. Top recommendation line
    for p in doc.paragraphs:
        if "Pro Appointments recommends" in p.text and "for an interview" in p.text:
            replace_paragraph_text(
                p,
                "Pro Appointments recommends {{candidate_first_name}} for an interview"
            )
            break

    # 2. Personal details rows
    for p in doc.paragraphs:
        for prefix, placeholder in field_appends.items():
            if p.text.startswith(prefix):
                if p.text.rstrip().endswith(":"):
                    append_text_to_paragraph(p, "\t" + placeholder)
                else:
                    append_text_to_paragraph(p, placeholder)
                break

    # 3. ITC / Criminal Record
    for i, p in enumerate(doc.paragraphs):
        if "ITC / CRIMINAL RECORD DISCLOSED" in p.text:
            for run in p.runs:
                if "None / Yes" in run.text:
                    run.text = run.text.replace("None / Yes", "{{itc_criminal_record}}")
                    break
            example_indices = []
            for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
                t = doc.paragraphs[j].text
                if t.startswith("If yes") or t.startswith("Candidate has declared"):
                    example_indices.append(j)
            if example_indices:
                replace_paragraph_text(doc.paragraphs[example_indices[0]], "{{itc_reason}}")
                for idx in example_indices[1:]:
                    replace_paragraph_text(doc.paragraphs[idx], "")
            break

    # 4. ACHIEVEMENTS & CERTIFICATIONS — find empty paragraph between heading tables
    paragraphs_between = find_paragraphs_between_headings(
        doc, "ACHIEVEMENTS & CERTIFICATIONS", "COMPUTER LITERACY"
    )
    placeholder_inserted = False
    for p in paragraphs_between:
        if not p.text.strip() and not placeholder_inserted:
            replace_paragraph_text(p, "{{achievements_block}}")
            placeholder_inserted = True
            break

    # 5. COMPUTER LITERACY & SKILLS — replace example bullet
    for p in doc.paragraphs:
        if "MS Office" in p.text and "Word" in p.text:
            replace_paragraph_text(p, "{{computer_skills_block}}")
            break

    # 6. EDUCATION (2 slots)
    edu_count = 0
    for p in doc.paragraphs:
        text = p.text
        if text.startswith("INSTITUTION\t:"):
            edu_count += 1
            if edu_count <= 2:
                append_text_to_paragraph(p, f"\t{{{{edu{edu_count}_institution}}}}")
        elif text.startswith("DATE\t:") and 1 <= edu_count <= 2:
            append_text_to_paragraph(p, f"\t{{{{edu{edu_count}_date}}}}")
        elif text.startswith("QUALIFICATION\t:") and 1 <= edu_count <= 2:
            append_text_to_paragraph(p, f"\t{{{{edu{edu_count}_qualification}}}}")
            if edu_count == 2:
                break

    # 7. EMPLOYMENT HISTORY (3 slots)
    # NOTE: order matters — check REASON FOR LEAVING before the in_duties branch,
    # otherwise it gets swallowed.
    job_count = 0
    in_duties = False
    duties_first_bullet_done = False

    for p in doc.paragraphs:
        text = p.text

        if text.startswith("COMPANY\t:"):
            job_count += 1
            in_duties = False
            duties_first_bullet_done = False
            if job_count <= 3:
                append_text_to_paragraph(p, f"\t{{{{job{job_count}_company}}}}")
        elif text.startswith("PERIOD OF EMPLOYMENT\t:") and 1 <= job_count <= 3:
            append_text_to_paragraph(p, f"\t{{{{job{job_count}_period}}}}")
        elif text.startswith("POSITION HELD\t:") and 1 <= job_count <= 3:
            append_text_to_paragraph(p, f"\t{{{{job{job_count}_position}}}}")
        elif text.startswith("REASON FOR LEAVING\t:") and 1 <= job_count <= 3:
            append_text_to_paragraph(p, f"\t{{{{job{job_count}_reason_for_leaving}}}}")
            in_duties = False
        elif text.strip() == "DUTIES" and 1 <= job_count <= 3:
            in_duties = True
        elif in_duties and 1 <= job_count <= 3:
            stripped = text.strip()
            if stripped == "M" and not duties_first_bullet_done:
                # First bullet under DUTIES — replace the "M" placeholder with our token
                replace_paragraph_text(p, f"{{{{job{job_count}_duties}}}}")
                duties_first_bullet_done = True
            # Leave the empty bullet paragraph as-is — the filler will handle expansion

    # 8. References — inject placeholders after each REFERENCE heading.
    # Structure in template:
    #   [REFERENCE 1 table] -> empty p -> empty p -> [REFERENCE 2 table] -> empty p -> "To Follow"
    # We replace the first empty paragraph after each heading with a reference block placeholder.

    # REFERENCE 1: replace the first empty paragraph between REFERENCE 1 and REFERENCE 2 headings
    ref1_paragraphs = find_paragraphs_between_headings(doc, "REFERENCE 1", "REFERENCE 2")
    for p in ref1_paragraphs:
        if not p.text.strip():
            replace_paragraph_text(p, "{{reference1_block}}")
            break

    # REFERENCE 2: walk after the REFERENCE 2 table; replace the first non-trivial paragraph
    # (either the empty one or the "To Follow" one — whichever comes first)
    body = doc.element.body
    found_ref2 = False
    ref2_done = False
    paragraphs_to_clear = []  # any "To Follow" text after we've placed our placeholder
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "tbl":
            table_text = "".join(child.itertext())
            if "REFERENCE 2" in table_text:
                found_ref2 = True
                continue
        elif tag == "p" and found_ref2:
            p_obj = Paragraph(child, doc)
            text = p_obj.text.strip()
            if not ref2_done:
                if not text:
                    replace_paragraph_text(p_obj, "{{reference2_block}}")
                    ref2_done = True
                elif text == "To Follow":
                    replace_paragraph_text(p_obj, "{{reference2_block}}")
                    ref2_done = True
            else:
                # We've placed the placeholder — clear any leftover "To Follow"
                if text == "To Follow":
                    paragraphs_to_clear.append(p_obj)

    for p in paragraphs_to_clear:
        replace_paragraph_text(p, "")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✅ Template with placeholders saved to {output_path}")


if __name__ == "__main__":
    input_doc = Path("/home/claude/Pro_Talent_CV_Master_Template_BLANK.docx")
    output_doc = Path("/home/claude/cv_automation/templates/pro_talent_template.docx")
    inject_placeholders(input_doc, output_doc)
