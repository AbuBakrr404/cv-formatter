"""
cv_parser.py
------------
Extracts raw text from CV files in PDF or DOCX format.
This is the first step in the pipeline — we get the text, then send it to Claude.
"""

from pathlib import Path
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_path: str | Path) -> str:
    """Extract text from a PDF using pdfplumber (better with layouts than PyPDF2)."""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_path: str | Path) -> str:
    """Extract text from a Word document, including tables."""
    doc = Document(file_path)
    text_parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Tables (many CVs use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def extract_cv_text(file_path: str | Path) -> str:
    """
    Main entry point — auto-detects file type and extracts text.
    Raises ValueError for unsupported file types.
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: {suffix}. "
            f"Supported types: .pdf, .docx"
        )


if __name__ == "__main__":
    # Quick test
    import sys
    if len(sys.argv) > 1:
        text = extract_cv_text(sys.argv[1])
        print(f"Extracted {len(text)} characters")
        print("---")
        print(text[:500])
